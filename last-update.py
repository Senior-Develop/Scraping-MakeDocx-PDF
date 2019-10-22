# -*- coding: utf-8 -*-
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
import os
import csv
import convertapi
import datetime
import pdfkit
import requests
import logging
import time
import re
from docx import Document
from datetime import date
import calendar
import shutil


logger = logging.getLogger("root")
logger.setLevel(logging.DEBUG)
# create console handler
ch = logging.StreamHandler()
ch.setLevel(logging.DEBUG)
logger.addHandler(ch)

API_KEY = "AIzaSyB9725AlCP1nHNvGFNvaWbkMN2ke_EWaAE"
# Backoff time sets how many minutes to wait between google pings when your API limit is hit
BACKOFF_TIME = 30
# Return Full Google Results? If True, full JSON results from Google are included in output
RETURN_FULL_RESULTS = False

output = "LetterInfo.csv"
output1 = "InvaildAddr.csv"
Addr_Val = True
# try:
#     os.remove(output)
# except OSError:
#     pass
if os.path.isfile(output) != True:
    line = ["BOOK_NO","FNAME", "LNAME", "ADDRESS1", "ADDRESS2", "CITY", "STATE", "ZIPCODE"]
    with open(output, 'w', newline='') as file1:
        writer = csv.writer(file1, delimiter=',')
        writer.writerow(line)

if os.path.isfile(output1) != True:
    line = ["BOOK_NO","FNAME", "LNAME", "ADDRESS1", "ADDRESS2", "CITY", "STATE", "ZIPCODE"]
    with open(output1, 'w', newline='') as file1:
        writer = csv.writer(file1, delimiter=',')
        writer.writerow(line)



def get_google_results(address, api_key=None, return_full_response=False):

    try:
        geocode_url = "https://maps.googleapis.com/maps/api/geocode/json?address={}".format(address)
        if api_key is not None:
            geocode_url = geocode_url + "&key={}".format(api_key)

        # Ping google for the reuslts:
        results = requests.get(geocode_url)
        # Results will be in JSON format - convert to dict using requests functionality
        results = results.json()

        # if there's no results or an error, return empty results.
        if len(results['results']) == 0:
            output = {
                "formatted_address": None,
                "latitude": None,
                "longitude": None,
                "accuracy": None,
                "google_place_id": None,
                "type": None,
                "postcode": None
            }
        else:
            answer = results['results'][0]
            output = {
                "formatted_address": answer.get('formatted_address')
            }
        output['status'] = results.get('status')

        if return_full_response is True:
            output['response'] = results

        return output
    except:
        pass


def address_validation(personal_info):
    global Addr_Val
    address = personal_info["ADDRESS1"] + "," + personal_info["CITY"] + "," + personal_info["STATE"] + "," + personal_info["ZIPCODE"]
    geocoded = False
    while geocoded is not True:
        # Geocode the address with google
        try:
            geocode_result = get_google_results(address, API_KEY, return_full_response=RETURN_FULL_RESULTS)
        except Exception as e:
            logger.exception(e)
            logger.error("Major error with {}".format(address))
            logger.error("Skipping!")
            geocoded = True

        # If we're over the API limit, backoff for a while and try again later.
        if geocode_result['status'] == 'OVER_QUERY_LIMIT':
            logger.info("Hit Query Limit! Backing off for a bit.")
            time.sleep(BACKOFF_TIME * 60)  # sleep for 30 minutes
            geocoded = False
        else:
            # If we're ok with API use, save the results
            # Note that the results might be empty / non-ok - log this
            if geocode_result['status'] != 'OK':
                logger.warning("Error geocoding {}: {}".format(address, geocode_result['status']))
            #logger.debug("Geocoded: {}: {}".format(address, geocode_result['status']))
            address_arr = {}
            # Append some other details:
            if geocode_result['status'] == 'OK':
                addrs = geocode_result['formatted_address'].split(",")
                if len(addrs) > 3:
                    personal_info['ADDRESS1'] = addrs[0]
                    personal_info['CITY'] = addrs[1]
                    stat = addrs[2].split(" ")
                    personal_info['STATE'] = stat[1]
                    personal_info['ZIPCODE'] = stat[2]
                else:
                    Addr_Val = False
                return personal_info
            geocoded = True



def get_driver():

    options = Options()
    options.add_experimental_option("excludeSwitches",
                                    ["ignore-certificate-errors", "safebrowsing-disable-download-protection",
                                     "safebrowsing-disable-auto-update", "disable-client-side-phishing-detection"])

    options.add_argument('--disable-infobars')
    options.add_argument('--disable-extensions')
    options.add_argument('--profile-directory=Default')
    options.add_argument("--incognito")
    options.add_argument("--disable-plugins-discovery")
    prefs = {'profile.default_content_setting_values.automatic_downloads': 1}
    options.add_experimental_option("prefs", prefs)
    #options.add_argument("--headless")
    driver = webdriver.Chrome('chromedriver', options=options)
    return driver


def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


def create_pdf(path):
    try:
        convertapi.api_secret = 'e29D5VW4X1vXpNWx'
        result = convertapi.convert('pdf', {'File': path})
        pdf_path = path[:-5] + ".pdf"
        # save to file
        result.file.save(pdf_path)
    except:
        pass



def create_docx(personal_info):
    savepath = os.getcwd() + "\\" + personal_info["BOOK_NO"]
    try:
        if not os.path.exists(savepath):
            os.makedirs(savepath)
        savefile = savepath + '\\Letter.docx'
        if os.path.isfile(savefile) != True:
            regex1 = re.compile("FIRST_NAME")
            replace1 = personal_info["FNAME"]
            filename = "LetterTemp.docx"
            doc = Document(filename)
            docx_replace_regex(doc, regex1, replace1)

            my_date = date.today()
            week_day = calendar.day_name[my_date.weekday()]
            month = calendar.month_name[my_date.month]
            Current_Time = week_day + ", " + month + " " + str(my_date.day) + ", " + str(my_date.year)
            regex1 = re.compile("Current_Time")
            replace1 = Current_Time
            docx_replace_regex(doc, regex1, replace1)
            doc.save(savefile)
        create_pdf(savefile)
        savefile = savepath + '\\LetterEnvelope.docx'
        if os.path.isfile(savefile) != True:
            regex1 = re.compile("FNAME")
            replace1 = personal_info["FNAME"]
            filename = "LetterEnvelopeTemp.docx"
            doc = Document(filename)
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("LNAME")
            replace1 = personal_info["LNAME"]
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("ADDRESS1")
            replace1 = personal_info['ADDRESS1']
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("ADDRESS2")
            replace1 = personal_info['ADDRESS2']
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("CITY")
            replace1 = personal_info['CITY'][1:]
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("STATE")
            replace1 = personal_info['STATE']
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("ZIPCODE")
            replace1 = personal_info['ZIPCODE']
            docx_replace_regex(doc, regex1, replace1)

            doc.save(savefile)
        create_pdf(savefile)
    except:
        pass



def main():

    BOOK = "NO.txt"
    global Addr_Val
    if os.path.isfile(BOOK):
        with open(BOOK, 'r') as filehandle:
            BOOK_NO = filehandle.readline()
            if not BOOK_NO:
                BOOK_NO = input("Please  Input BOOKING NO :")
    else:
        BOOK_NO = input("Please  Input BOOKING NO :")
    driver = get_driver()
    try:

        print(datetime.datetime.now())

        url = "https://apps.co.lubbock.tx.us/jailrosters/activejail.aspx"
        driver.get(url)

        table = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.ID, 'gridaj')))
        BOOKING = table.find_element_by_tag_name("td")
        BOOKING.click()
        table = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.ID, 'gridaj')))
        BOOKING = table.find_element_by_tag_name("td")
        BOOKING.click()
        Next_check = 0
        BOOK_check = True
        MAX_check = 0

        while BOOK_check == True:
            table = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.ID, 'gridaj')))
            BOOK_ARR = []
            trs = table.find_elements_by_tag_name("tr")
            for i, tr in enumerate(trs):
                if i != 0:
                    BOOkN = tr.find_element_by_tag_name("td")
                    BOOK_ARR.append(BOOkN.text)
            if MAX_check == 0:
                BOOK_MAX = BOOK_ARR[0]
                MAX_check += 1
            sheets = table.find_elements_by_tag_name("input")
            sheet_len = len(sheets)
            idx = 0
            while idx < sheet_len:
                person = {}
                table = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.ID, 'gridaj')))
                if int(BOOK_ARR[idx]) > int(BOOK_NO):
                    sheets = table.find_elements_by_tag_name("input")
                    main_window_handle = None
                    while not main_window_handle:
                        main_window_handle = driver.current_window_handle
                    sheets[idx].click()
                    time.sleep(1)
                    signin_window_handle = None
                    while not signin_window_handle:
                        for handle in driver.window_handles:
                            if handle != main_window_handle:
                                signin_window_handle = handle
                                break
                    driver.switch_to.window(signin_window_handle)
                    time.sleep(1)
                    driver.switch_to.frame(driver.find_element_by_tag_name("frame"))
                    time.sleep(1)
                    tables = driver.find_elements_by_tag_name("table")
                    if len(tables) != 10:
                        addr = driver.find_element_by_id("addr")
                        address = addr.text
                        if ("HOMELESS" not in address) and (address != "") :

                            person["BOOK_NO"] = BOOK_ARR[idx]
                            Name_Label = driver.find_element_by_id("Label1")
                            Name = Name_Label.text
                            LName = Name[:Name.index(",")]
                            Names = Name[Name.index(",") + 2 :]
                            Name_a = Names.split(" ")
                            FName = Name_a[0]
                            FIRST_NAME = FName.title()
                            LAST_NAME = LName.title()
                            person["FNAME"] = FIRST_NAME
                            person["LNAME"] = LAST_NAME

                            if "#" in address:
                                person["ADDRESS1"] = address[:address.index("#")]
                                person["ADDRESS2"] = address[address.index("#") :]
                            elif "APT" in address:
                                person["ADDRESS1"] = address[:address.index("APT")]
                                person["ADDRESS2"] = address[address.index("APT"):]
                            elif "SUITE" in address:
                                person["ADDRESS1"] = address[:address.index("SUITE")]
                                person["ADDRESS2"] = address[address.index("SUITE"):]
                            else:
                                person["ADDRESS1"] = address
                                person["ADDRESS2"] = ""

                            Citys = driver.find_element_by_id("citystzip")
                            City_Zip = Citys.text
                            citys_arr = City_Zip.split(" ")

                            for idk , city in enumerate(citys_arr):
                                if idk == 0:
                                    person["CITY"] = city[:-1]
                                elif idk == 1:
                                    person["STATE"] = city
                                elif idk == 2:
                                    person["ZIPCODE"] = city

                            updateperson = address_validation(person)
                            #create_docx(updateperson)
                            if Addr_Val == True:
                                dir_path = os.getcwd()
                                directory = dir_path + "\\" + BOOK_ARR[idx]
                                directory1 = dir_path + "\\allpdfs"

                                if not os.path.exists(directory):
                                    os.makedirs(directory)

                                if not os.path.exists(directory1):
                                    os.makedirs(directory1)

                                page_html = driver.page_source
                                page_result = page_html.replace("../lsoimages","http://apps.co.lubbock.tx.us/lsoimages")
                                export_pdf = directory + "\\" + BOOK_ARR[idx] + ".pdf"
                                export_pdf1 = directory1 + "\\" + BOOK_ARR[idx] + ".pdf"
                                pdfkit.from_string(page_result, export_pdf)

                                if os.path.isfile(export_pdf1) != True:
                                    shutil.copyfile(export_pdf, export_pdf1)

                                keys = updateperson.keys()
                                with open(output, 'a',newline='') as output_file:
                                    dict_writer = csv.DictWriter(output_file, keys)
                                    dict_writer.writerow(updateperson)

                            else:
                                Addr_Val = True
                                print("Invalid Address : " + person["FNAME"] + ", " + person["LNAME"] + "@ " + person["ADDRESS1"] + ", " + person["ADDRESS2"] + ", " + person["CITY"] + ", " + person["STATE"] + ", " + person["ZIPCODE"])
                                keys = person.keys()
                                with open(output1, 'a',newline='') as output_file:
                                    dict_writer = csv.DictWriter(output_file, keys)
                                    dict_writer.writerow(updateperson)

                    else:
                        BOOK_MAX = BOOK_ARR[idx]
                    time.sleep(1)

                    driver.switch_to.default_content()
                    time.sleep(1)
                    driver.switch_to.window(main_window_handle)
                else:
                    BOOK_check = False
                    break
                idx += 1
            if Next_check == 0:
                Next = driver.find_element_by_xpath("//*[@id='gridaj']/tbody/tr[12]/td/a")
                Next_check += 1
            else:
                Next = driver.find_element_by_xpath("//*[@id='gridaj']/tbody/tr[12]/td/a[2]")
            Next.click()
            with open(BOOK, 'w') as the_file:
                the_file.write(BOOK_MAX)

    except:
        print("loading page error")
        pass
    driver.quit()
    print("---------processing end-----------")
    print(datetime.datetime.now())
    time.sleep(60 * 30)

if __name__ == "__main__":
    while True:
        main()
