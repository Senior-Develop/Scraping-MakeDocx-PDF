# -*- coding: utf-8 -*-
import os
import csv
import re
from docx import Document
from datetime import date
import calendar
import os
from docxcompose.composer import Composer
from docx import Document as Document_compose
from PyPDF2 import PdfFileMerger
from time import gmtime, strftime



BOOK = "UPDATE.txt"
if os.path.isfile(BOOK):
    with open(BOOK, 'r') as filehandle:
        BOOK_NO = filehandle.readline()
        if not BOOK_NO:
            BOOK_NO = input("Please  Input BOOKING NO :")
else:
    BOOK_NO = input("Please  Input BOOKING NO :")

name = strftime("%Y%m%d%H%M", gmtime())
directory = os.getcwd() + "\\print"
if not os.path.exists(directory):
    os.makedirs(directory)


def combine_all_letter(filename_master,files_list):

    savepath = os.getcwd() + "\\print\\" + name + "_combine.docx"
    number_of_sections=len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        if os.path.isfile(files_list[i]) == True:
            doc_temp = Document_compose(files_list[i])
            composer.append(doc_temp)
    composer.save(savepath)


def combine_all_env(filename_master,files_list):

    savepath = os.getcwd() + "\\print\\" + name + "_combined_envelope.docx"
    number_of_sections=len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        if os.path.isfile(files_list[i]) == True:
            doc_temp = Document_compose(files_list[i])
            composer.append(doc_temp)
    composer.save(savepath)



def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


def create_docx(personal_info):

    savepath = os.getcwd() + "\\" + personal_info["BOOK_NO"]
    if not os.path.exists(savepath):
        os.makedirs(savepath)
    regex1 = re.compile("FIRST_NAME")
    replace1 = personal_info["FNAME"]
    filename = "LetterTemp.docx"
    doc = Document(filename)
    docx_replace_regex(doc, regex1, replace1)

    my_date = date.today()
    week_day = calendar.day_name[my_date.weekday()]
    month = calendar.month_name[my_date.month]
    Current_Time = week_day + ", " + month + " " + str(my_date.day) + ", " + str(my_date.year)
    regex1 = re.compile("Current_Time")
    replace1 = Current_Time
    docx_replace_regex(doc, regex1, replace1)
    doc.save(savepath + '\\Letter.docx')

    regex1 = re.compile("FNAME")
    replace1 = personal_info["FNAME"]
    filename = "LetterEnvelopeTemp.docx"
    doc = Document(filename)
    docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile("LNAME")
    replace1 = personal_info["LNAME"]
    docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile("ADDRESS1")
    replace1 = personal_info['ADDRESS1']
    docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile("ADDRESS2")
    replace1 = personal_info['ADDRESS2']
    docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile("CITY")
    replace1 = personal_info['CITY'][1:]
    docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile("STATE")
    replace1 = personal_info['STATE']
    docx_replace_regex(doc, regex1, replace1)

    regex1 = re.compile("ZIPCODE")
    replace1 = personal_info['ZIPCODE']
    docx_replace_regex(doc, regex1, replace1)

    doc.save(savepath + '\\LetterEnvelope.docx')


def main():
    BOOK_MAX = "0"
    BO_OK = []
    with open('LetterInfo.csv', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            if int(row["BOOK_NO"]) > int(BOOK_NO):
                personal_info = {
                    "BOOK_NO" : row["BOOK_NO"],
                    "FNAME" : row["FNAME"],
                    "LNAME" :row["LNAME"],
                    "ADDRESS1" : row["ADDRESS1"],
                    "ADDRESS2" : row["ADDRESS2"],
                    "CITY" : row["CITY"],
                    "STATE" : row["STATE"],
                    "ZIPCODE" : row["ZIPCODE"]
                }
                create_docx(personal_info)
                BO_OK.append(personal_info["BOOK_NO"])

    Path = os.getcwd()
    files = []
    if len(BO_OK) > 0:
        BOOK_MAX = BO_OK[0]
        for OK in BO_OK:
            files.append(Path + "\\" + OK + "\\Letter.docx")

        first_file = files.pop(0)
        combine_all_letter(first_file, files)

        files = []
        for OK in BO_OK:
            files.append(Path + "\\" + OK + "\\LetterEnvelope.docx")

        first_file = files.pop(0)
        combine_all_env(first_file, files)

        savepath = os.getcwd() + "\\print\\" + name + "_combine.pdf"

        pdfs = []
        for OK in BO_OK:
            pdfs.append(Path + "\\" + OK + "\\" + OK + ".pdf")
        merger = PdfFileMerger()

        for pdf in pdfs:
            if os.path.isfile(pdf) == True:
                merger.append(pdf)

        merger.write(savepath)
        merger.close()
        with open(BOOK, 'w') as the_file:
            the_file.write(BOOK_MAX)

if __name__ == "__main__":

    main()
    # schedule.every().day.at("08:00").do(main)
    #
    # schedule.every().day.at("16:00").do(main)
    # while True:
    #     schedule.run_pending()
    #     time.sleep(60)