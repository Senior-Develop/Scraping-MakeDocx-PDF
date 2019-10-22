# -*- coding: utf-8 -*-
import os
import csv
import re
from docx import Document
from datetime import date
import calendar
import os
from PyPDF2 import PdfFileMerger
from time import gmtime, strftime
import convertapi
import datetime
import schedule
import time

BOOK = "UPDATE.txt"
if os.path.isfile(BOOK):
    with open(BOOK, 'r') as filehandle:
        BOOK_NO = filehandle.readline()
        if not BOOK_NO:
            BOOK_NO = input("Please  Input BOOKING NO :")
else:
    BOOK_NO = input("Please  Input BOOKING NO :")

name = strftime("%Y%m%d%H%M", gmtime())
directory = os.getcwd() + "\\print"
if not os.path.exists(directory):
    os.makedirs(directory)



def create_pdf(path):
    try:
        convertapi.api_secret = 'e29D5VW4X1vXpNWx'
        result = convertapi.convert('pdf', {'File': path})
        pdf_path = path[:-5] + ".pdf"
        # save to file
        result.file.save(pdf_path)
    except:
        pass


def merge_pdf(pdfs, savepath):

    try:
        merger = PdfFileMerger(strict=False)

        for pdf in pdfs:
            if os.path.isfile(pdf) == True:
                merger.append(pdf)

        merger.write(savepath)
        merger.close()
    except:
        pass

# def combine_all_letter(files_list):
#
#     savepath = os.getcwd() + "\\print\\" + name + "_combine.docx"
#
#     merged_document = Document()
#
#     for index, file in enumerate(files_list):
#         if os.path.isfile(file) == True:
#             sub_doc = Document(file)
#
#             # Don't add a page break if you've reached the last file.
#             if index < len(files_list)-1:
#                sub_doc.add_page_break()
#
#             for element in sub_doc.element.body:
#                 merged_document.element.body.append(element)
#
#     merged_document.save(savepath)
#
#
# def combine_all_env(files_list):
#
#     savepath = os.getcwd() + "\\print\\" + name + "_combined_envelope.docx"
#
#     merged_document = Document()
#
#     for index, file in enumerate(files_list):
#         if os.path.isfile(file) == True:
#             sub_doc = Document(file)
#
#             # Don't add a page break if you've reached the last file.
#             if index < len(files_list)-1:
#                sub_doc.add_page_break()
#
#             for element in sub_doc.element.body:
#                 merged_document.element.body.append(element)
#
#     merged_document.save(savepath)

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


def create_docx(personal_info):

    savepath = os.getcwd() + "\\" + personal_info["BOOK_NO"]
    try:
        if not os.path.exists(savepath):
            os.makedirs(savepath)
        savefile = savepath + '\\Letter.docx'
        if os.path.isfile(savefile) != True:
            print("create docx: " + personal_info["BOOK_NO"] + "\n")
            regex1 = re.compile("FIRST_NAME")
            replace1 = personal_info["FNAME"]
            filename = "LetterTemp.docx"
            doc = Document(filename)
            docx_replace_regex(doc, regex1, replace1)

            my_date = date.today()
            week_day = calendar.day_name[my_date.weekday()]
            month = calendar.month_name[my_date.month]
            Current_Time = week_day + ", " + month + " " + str(my_date.day) + ", " + str(my_date.year)
            regex1 = re.compile("Current_Time")
            replace1 = Current_Time
            docx_replace_regex(doc, regex1, replace1)
            doc.save(savefile)
        create_pdf(savefile)
        savefile = savepath + '\\LetterEnvelope.docx'
        if os.path.isfile(savefile) != True:
            regex1 = re.compile("FNAME")
            replace1 = personal_info["FNAME"]
            filename = "LetterEnvelopeTemp.docx"
            doc = Document(filename)
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("LNAME")
            replace1 = personal_info["LNAME"]
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("ADDRESS1")
            replace1 = personal_info['ADDRESS1']
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("ADDRESS2")
            replace1 = personal_info['ADDRESS2']
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("CITY")
            replace1 = personal_info['CITY'][1:]
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("STATE")
            replace1 = personal_info['STATE']
            docx_replace_regex(doc, regex1, replace1)

            regex1 = re.compile("ZIPCODE")
            replace1 = personal_info['ZIPCODE']
            docx_replace_regex(doc, regex1, replace1)

            doc.save(savefile)
        create_pdf(savefile)
    except:
        pass

def main():
    BOOK_MAX = "0"
    BO_OK = []
    print((datetime.datetime.now()))
    print("\n")
    print("-------------------------\n")
    with open('LetterInfo.csv', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            if int(row["BOOK_NO"]) > int(BOOK_NO):
                personal_info = {
                    "BOOK_NO" : row["BOOK_NO"],
                    "FNAME" : row["FNAME"],
                    "LNAME" :row["LNAME"],
                    "ADDRESS1" : row["ADDRESS1"],
                    "ADDRESS2" : row["ADDRESS2"],
                    "CITY" : row["CITY"],
                    "STATE" : row["STATE"],
                    "ZIPCODE" : row["ZIPCODE"]
                }
                if personal_info["BOOK_NO"] not in BO_OK:

                    create_docx(personal_info)
                    BO_OK.append(personal_info["BOOK_NO"])
                    BO_OK = list(set(BO_OK))

    Path = os.getcwd()
    files = []
    if len(BO_OK) > 0:
        print("create merge PDF\n")

        BOOK_MAX = BO_OK[0]
        for OK in BO_OK:
            files.append(Path + "\\" + OK + "\\Letter.pdf")

        savepath = os.getcwd() + "\\print\\" + name + "_letter.pdf"
        merge_pdf(files,savepath)
        #first_file = files.pop(0)
        #combine_all_letter(files)

        files = []
        for OK in BO_OK:
            files.append(Path + "\\" + OK + "\\LetterEnvelope.pdf")
        savepath = os.getcwd() + "\\print\\" + name + "_LetterEnvelope.pdf"
        merge_pdf(files, savepath)
        #first_file = files.pop(0)
        #combine_all_env(files)

        savepath = os.getcwd() + "\\print\\" + name + "_combine.pdf"

        pdfs = []
        for OK in BO_OK:
            pdfs.append(Path + "\\" + OK + "\\" + OK + ".pdf")

        merge_pdf(pdfs, savepath)

        with open(BOOK, 'w') as the_file:
            the_file.write(BOOK_MAX)

    print((datetime.datetime.now()))

if __name__ == "__main__":

    print((datetime.datetime.now()))
    print("\n")
    print("-------------------------\n")
    main()
    print((datetime.datetime.now()))
    # schedule.every().day.at("08:00").do(main)
    #
    # schedule.every().day.at("16:00").do(main)
    # while True:
    #     schedule.run_pending()
    #     time.sleep(60)